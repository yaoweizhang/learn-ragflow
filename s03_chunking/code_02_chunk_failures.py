#!/usr/bin/env python3
"""
s03 / unit 02 — 把 unit 01 的 chunker 跑在真实样本上,展示 3 类典型失败:

(a) 表格被切碎 — 整机规格表被硬切到 mid-row,每个 chunk 只看到半张表
(b) 父子块缺失 — "Q3 营收多少" 命中节标题 chunk,数字本身在 DOCX 表里(被 s02 loader 丢了)
(c) 跨段引用断裂 — "第二节 主要财务数据" header 单独成 10-char chunk,语义为零

每个 demo 都打印 before/after 片段,让失败肉眼可见。

对照 ragflow 的解法: XGBoost _concat_downward (父块) + naive_merge (token 子块)
+ hierarchical_merge (层级) + attach_media_context (表格上下文回填)。

运行: python s03_chunking/code_02_chunk_failures.py
需要: 同 unit 01
"""
import importlib.util
import sys
from pathlib import Path

# 复用 unit 01 的 chunker——章节内 import 同 unit 是允许的(为什么要拆 unit)。
_UNIT01_PATH = Path(__file__).resolve().parents[0] / "code_01_basic_chunk.py"
_spec = importlib.util.spec_from_file_location("s03_unit01_basic_chunk", _UNIT01_PATH)
_mod = importlib.util.module_from_spec(_spec)
sys.modules[_spec.name] = _mod
_spec.loader.exec_module(_mod)
load_pdf = _mod.load_pdf
load_docx = _mod.load_docx
chunk_by_paragraph = _mod.chunk_by_paragraph
split_long_paragraph = _mod.split_long_paragraph

WORKDIR = Path(__file__).resolve().parents[1]
SAMPLES = WORKDIR / "samples"


def _hr(title: str) -> None:
    print()
    print("=" * 72)
    print(title)
    print("=" * 72)


def demo_table_split() -> None:
    """(a) 表格被切碎: 整机规格表 562 字符 → 硬切到 mid-row."""
    _hr("[a] 表格被切碎 — 整机规格表 hard-cut")
    pdf = load_pdf(SAMPLES / "server_whitepaper.pdf")
    spec_doc = next(d for d in pdf if "三、整机规格" in d["text"])
    print(f"BEFORE (整段 {len(spec_doc['text'])} 字符):")
    print(spec_doc["text"][:240], "...")
    print()

    pieces = split_long_paragraph(spec_doc["text"], max_chars=500)
    print(f"AFTER  (cap=500 → 切成 {len(pieces)} 块):")
    for i, p in enumerate(pieces):
        last_line = p.strip().splitlines()[-1] if p.strip() else "<空>"
        print(f"  chunk {i} len={len(p):>4} | 末行: {last_line[:60]}")
    print()
    last0 = pieces[0].strip().splitlines()[-1] if pieces else "<空>"
    print(f"→ 失败点: 第 0 块末行停在 '{last0}' — 表的列对齐被腰斩,")
    print(f"          第 1 块从 'NAND' 段继续,失去 '组件 / 规格 / 说明' 列对齐语义。")
    print("→  ragflow 修法: deepdoc/parser/pdf_parser.py 用 XGBoost 30 特征识别表格 layout_type=table,")
    print("    把整张表当 parent 块;rag/nlp/__init__.py:attach_media_context 给表回填前后文本 context。")


def demo_parent_child() -> None:
    """(b) 父子块缺失: 节标题单独成 chunk, 数据本体在 DOCX 表里被 s02 丢了."""
    _hr("[b] 父子块缺失 — 节标题 chunk 与数据本体分离")
    chunks = chunk_by_paragraph(
        load_pdf(SAMPLES / "server_whitepaper.pdf") + load_docx(SAMPLES / "disclosure.docx")
    )
    # 找出"分季度财务数据"这个 header 单独成 chunk 的实例
    header_chunks = [c for c in chunks if c["text"].strip() in {"第四节 分季度财务数据", "第二节 主要财务数据"}]
    print("BEFORE (用户问 'Q3 营收多少'):")
    for c in header_chunks:
        print(f"  召回命中 chunk_id={c['chunk_id']} len={len(c['text'])} text='{c['text']}'")
    print()

    # 实际季度数据在 DOCX tables 里——s02 unit 01 的 loader 不读 tables,这里直接展示原文
    from docx import Document
    doc = Document(SAMPLES / "disclosure.docx")
    print(f"AFTER  (季度数据本体在 DOCX tables,共 {len(doc.tables)} 张表):")
    for i, tbl in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
        if rows and any("Q3" in cell or "三季度" in cell or "9 月" in cell for cell in rows[0]):
            print(f"  Table {i} (含 Q3 数据):")
            for row in rows[:5]:
                print(f"    {row}")
            break
    else:
        print("  (本样本 DOCX 表未含 Q3 字面量; 但所有季度数字都只在 tables 里,chunker 看不到)")
    print()
    print("→ 失败点: 检索命中 10-char 的 '第四节 分季度财务数据' header,")
    print("          LLM 收到后只能反问 '请问您指哪个季度'——数据本体在 s02 loader 丢弃的 DOCX tables 里。")
    print("→  ragflow 修法: deepdoc/parser/docx_parser.py 把 tables 也读出来当 parent 块,")
    print("    hierarchical_merge 按 '第 X 节' 正则建父子树;召回时返回 parent 全文。")


def demo_cross_ref() -> None:
    """(c) 跨段引用断裂: '如下' / '见上表' 这种指代词单独成 chunk 后无意义."""
    _hr("[c] 跨段引用断裂 — 指代词单独成 chunk")
    chunks = chunk_by_paragraph(
        load_pdf(SAMPLES / "server_whitepaper.pdf") + load_docx(SAMPLES / "disclosure.docx")
    )
    # 找出过短(< 30 字符)的 header-only chunk
    print("BEFORE (过短的 header-only chunks,语义为零):")
    short = [c for c in chunks if 0 < len(c["text"]) < 30]
    for c in short:
        print(f"  id={c['chunk_id']} len={len(c['text']):>3} | '{c['text']}'")
    print()
    print(f"AFTER  (DOCX 原文: '按业务板块划分,2024 年公司收入结构如下:...'):")
    # 把真实的指代上下文找出来
    docs = load_docx(SAMPLES / "disclosure.docx")
    ref_doc = next(d for d in docs if "收入结构如下" in d["text"])
    print(f"  '{ref_doc['text'][:200]}...'")
    print(f"  (后续'智能算力 ...' 等具体数字段都另起段落,被 chunker 切成独立 chunk,")
    print(f"   '如下' 这个指代词所在的 chunk 完全不知道它指的是谁)")
    print()
    print("→ 失败点: '收入结构如下' 单独成 chunk;命中它的话 LLM 看到的是不完整句子。")
    print("→  ragflow 修法: hierarchical_merge 按标题级别建层级,把 '如下' 所在 chunk 与后续")
    print("    数字 chunk 绑成同一节;召回命中其一就附带整节。")


def main() -> None:
    demo_table_split()
    demo_parent_child()
    demo_cross_ref()
    print()
    print("→ 三类失败都不是 unit 01 的 chunker 单独造成的——是 chunker + s02 loader + 段落切分")
    print("  的累积效应。ragflow 用 4 层流水线 (XGBoost 父块 + token 子块 + 层级合并 + 媒体回填)")
    print("  逐层修。")


if __name__ == "__main__":
    main()