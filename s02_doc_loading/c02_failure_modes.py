#!/usr/bin/env python3
"""
s02 / unit 02 — 失败模式：把 unit 01 的 loader 跑在真实样本上,
展示 (a) PDF 多栏错位 (b) DOCX 表格被吞。

对照 ragflow 的 deepdoc/parser/：用 XGBoost 版面分析 + table-aware 解析修这些问题。

运行: python s02_doc_loading/code_02_failure_modes.py
需要: 同 unit 01 + samples/{server_whitepaper.pdf,disclosure.docx}
"""
import importlib.util
import sys
from pathlib import Path

# 复用 unit 01 的 loader（章节内 import 是允许的——这就是为什么要拆 unit）。
# 目录以数字开头，普通 `import` 报 SyntaxError，改用 importlib 加载。
_UNIT01_PATH = Path(__file__).resolve().parents[0] / "code_01_basic_load.py"
_spec = importlib.util.spec_from_file_location("s02_unit01_basic_load", _UNIT01_PATH)
_mod = importlib.util.module_from_spec(_spec)
sys.modules[_spec.name] = _mod
_spec.loader.exec_module(_mod)
load_pdf = _mod.load_pdf
load_docx = _mod.load_docx

WORKDIR = Path(__file__).resolve().parents[1]
SAMPLES = WORKDIR / "samples"


def show_pdf_failure() -> None:
    pdf = load_pdf(SAMPLES / "server_whitepaper.pdf")
    print(f"[PDF] {len(pdf)} 页抽出的段落 (page, len, first 60 字):")
    for seg in pdf:
        print(f"  page={seg['page']:>2} len={len(seg['text']):>4} | {seg['text'][:60].replace(chr(10), ' ')}")


def show_docx_table_loss() -> None:
    from docx import Document
    path = SAMPLES / "disclosure.docx"
    doc = Document(path)
    para_count = sum(1 for p in doc.paragraphs if p.text.strip())
    table_count = len(doc.tables)
    table_text_len = sum(
        len(cell.text)
        for tbl in doc.tables
        for row in tbl.rows
        for cell in row.cells
    )
    print(f"\n[DOCX] paragraphs(非空)={para_count}, tables={table_count}, 表格内总字符={table_text_len}")
    print(f"  → unit 01 的 load_docx 只读 paragraphs，丢失 {table_text_len} 字符（{table_count} 张表）")


def main() -> None:
    show_pdf_failure()
    show_docx_table_loss()
    print("\n→ ragflow 的解法: deepdoc/parser/pdf_parser.py 用 XGBoost 版面分析;")
    print("  deepdoc/parser/docx_parser.py 同时遍历 paragraphs + tables")


if __name__ == "__main__":
    main()
