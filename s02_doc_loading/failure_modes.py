#!/usr/bin/env python3
"""
s02 / unit 02 — 失败模式：把 unit 01 的 loader 跑在真实样本上,
展示 (a) PDF 多栏错位 (b) DOCX 表格被吞。

对照 RAGFlow 的 deepdoc/parser/：用 XGBoost 版面分析 + table-aware 解析修这些问题。

运行: python s02_doc_loading/failure_modes.py
需要: 同 unit 01 + samples/{server_whitepaper.pdf,disclosure.docx}

术语速览 (本文件首次出现):
- 版面分析 (layout analysis): 把页面切分为标题/正文/表格/图片等区域并标注类型
- XGBoost: 梯度提升树模型,RAGFlow 用来训练版面分类器
- 多栏错位 (multi-column reflow): PDF 双栏排版被按行扫,左右栏文本被串成一行
- importlib.util.spec_from_file_location: 用文件路径加载 Python 模块的 API
- DOCX 表格: Word 里的 table 对象,与 paragraph 平级;普通 loader 只读 paragraph 会丢表
"""
import importlib.util
import sys
from pathlib import Path

# 复用 unit 01 的 loader（章节内 import 是允许的——这就是为什么要拆 unit）。
# 目录以数字开头，普通 `import` 报 SyntaxError，改用 importlib 加载。
_UNIT01_PATH = Path(__file__).resolve().parents[0] / "basic_load.py"
_spec = importlib.util.spec_from_file_location("s02_unit01_basic_load", _UNIT01_PATH)
_mod = importlib.util.module_from_spec(_spec)
sys.modules[_spec.name] = _mod
_spec.loader.exec_module(_mod)
load_pdf = _mod.load_pdf
load_docx = _mod.load_docx

WORKDIR = Path(__file__).resolve().parents[1]
SAMPLES = WORKDIR / "samples"


def show_pdf_failure() -> None:
    pdf = load_pdf(SAMPLES / "server_whitepaper.pdf")
    print(f"[PDF] {len(pdf)} 页抽出的段落 (page, len, first 60 字):")
    for seg in pdf:
        print(f"  page={seg['page']:>2} len={len(seg['text']):>4} | {seg['text'][:60].replace(chr(10), ' ')}")


def show_docx_table_loss() -> None:
    from docx import Document
    path = SAMPLES / "disclosure.docx"
    doc = Document(path)
    para_count = sum(1 for p in doc.paragraphs if p.text.strip())
    table_count = len(doc.tables)
    table_text_len = sum(
        len(cell.text)
        for tbl in doc.tables
        for row in tbl.rows
        for cell in row.cells
    )
    print(f"\n[DOCX] paragraphs(非空)={para_count}, tables={table_count}, 表格内总字符={table_text_len}")
    print(f"  → c01 的 load_docx 只读 paragraphs，丢失 {table_text_len} 字符（{table_count} 张表）")


def main() -> None:
    show_pdf_failure()
    show_docx_table_loss()
    print("\n→ RAGFlow 的解法: deepdoc/parser/pdf_parser.py 用 XGBoost 版面分析;")
    print("  deepdoc/parser/docx_parser.py 同时遍历 paragraphs + tables")


if __name__ == "__main__":
    main()
